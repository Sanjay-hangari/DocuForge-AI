import streamlit as st
from docx import Document
from openai import OpenAI
from datetime import timedelta
from docx.shared import Pt
import os, json, tempfile, re

MODEL = "openai/gpt-4o-mini"
TEMPLATE_PATH = "Templetes/Template.docx"
WORK_HOURS_PER_DAY = 8
API_key = st.secrets["OPENROUTER_API_KEY"]
client = OpenAI(
    base_url="https://openrouter.ai/api/v1",
    api_key = API_key
)

# ================= UI =================
st.set_page_config("DocuForge AI", layout="centered")
st.title("📄✨ DocuForge AI")

overview = st.text_area("Project Overview", height=120)
exec_summary = st.text_area("Executive Summary", height=120)

col1, col2 = st.columns(2)
with col1:
    start_date = st.date_input("Start Date")
with col2:
    end_date = st.date_input("End Date")

# ================= PROMPT =================
def build_prompt(overview, exec_summary):
    return f"""
You are a senior project manager.

Tasks:
1. Rewrite and expand project overview professionally.
2. Rewrite and expand executive summary professionally.

3. Generate DETAILED IN-SCOPE:
   - Each item must be a full detailed professional sentence
   - Add missing logical scope items
   - Minimum 6 items

4. Generate OUT OF SCOPE.
5. Generate ASSUMPTIONS.
6. Generate RISKS with MITIGATION.

7. Generate HIGH LEVEL PLAN:
   - 6 to 10 main tasks
   - Each task must include effort_days
   - Effort reflects complexity

Return ONLY JSON.

{{
 "overview": "...",
 "executive_summary": "...",
 "in_scope": ["..."],
 "out_scope": ["..."],
 "assumptions": ["..."],
 "tasks": [
   {{ "name": "...", "effort_days": 0 }}
 ],
 "risks": [
   {{ "risk": "...", "mitigation": "..." }}
 ]
}}

Overview:
{overview}

Executive Summary:
{exec_summary}
"""

# ================= JSON PARSER =================
def extract_json(text):
    match = re.search(r"\{[\s\S]*\}", text)
    if not match:
        raise ValueError("No JSON found")

    json_text = match.group()
    json_text = re.sub(r",\s*}", "}", json_text)
    json_text = re.sub(r",\s*]", "]", json_text)

    return json.loads(json_text)

def call_llm(prompt):
    res = client.chat.completions.create(
        model=MODEL,
        messages=[
            {"role": "system", "content": "Return valid JSON only."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.3,
        max_tokens=2000
    )
    return extract_json(res.choices[0].message.content)

# ================= WORKING DAY LOGIC =================
def is_working_day(d):
    return d.weekday() < 5  # Mon-Fri

def working_days_between(start, end):
    count = 0
    d = start
    while d <= end:
        if is_working_day(d):
            count += 1
        d += timedelta(days=1)
    return count

def add_working_days(start, days):
    current = start
    added = 0
    while added < days:
        if is_working_day(current):
            added += 1
        if added < days:
            current += timedelta(days=1)
    return current

# ================= SCHEDULER =================
def build_schedule(tasks, start, end):
    total_working_days = working_days_between(start, end)
    total_effort = sum(t["effort_days"] for t in tasks)

    # SCALE to use full timeline
    scale = total_working_days / total_effort
    for t in tasks:
        t["effort_days"] = max(1, round(t["effort_days"] * scale))

    current = start
    schedule = []

    for t in tasks:
        task_start = current
        task_end = add_working_days(task_start, t["effort_days"])
        hours = t["effort_days"] * WORK_HOURS_PER_DAY

        schedule.append({
            "task": t["name"],
            "start": task_start.strftime("%d-%m-%Y"),
            "end": task_end.strftime("%d-%m-%Y"),
            "hours": hours
        })

        current = task_end + timedelta(days=1)

    return schedule

# ================= TABLE BORDER =================
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_border(table):
    tbl = table._tbl
    borders = OxmlElement('w:tblBorders')
    for edge in ('top','left','bottom','right','insideH','insideV'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), '4')
        elem.set(qn('w:color'), '000000')
        borders.append(elem)
    tbl.tblPr.append(borders)

# ================= WORD =================
def generate_word(data, schedule, path):
    doc = Document(TEMPLATE_PATH)

    def heading(text):
        h = doc.add_heading(level=2)
        r = h.add_run(text)
        r.bold = True
        r.font.size = Pt(14)

    def bullets(items):
        for i in items:
            doc.add_paragraph("• " + str(i))

    heading("Project Overview")
    doc.add_paragraph(data["overview"])

    heading("Executive Summary")
    doc.add_paragraph(data["executive_summary"])

    heading("In Scope")
    bullets(data["in_scope"])

    heading("Out of Scope")
    bullets(data["out_scope"])

    heading("Assumptions")
    bullets(data["assumptions"])

    heading("High Level Plan")
    table = doc.add_table(rows=1, cols=4)
    set_table_border(table)

    hdr = table.rows[0].cells
    hdr[0].text = "Task"
    hdr[1].text = "Start"
    hdr[2].text = "End"
    hdr[3].text = "Effort (hrs)"

    for row in schedule:
        r = table.add_row().cells
        r[0].text = row["task"]
        r[1].text = row["start"]
        r[2].text = row["end"]
        r[3].text = str(row["hours"])

    heading("Risk & Mitigation")
    rtable = doc.add_table(rows=1, cols=2)
    set_table_border(rtable)

    rtable.rows[0].cells[0].text = "Risk"
    rtable.rows[0].cells[1].text = "Mitigation"

    for r in data["risks"]:
        row = rtable.add_row().cells
        row[0].text = r["risk"]
        row[1].text = r["mitigation"]

    doc.save(path)

# ================= MAIN =================
if st.button("Generate Report"):

    with st.spinner("Generating report..."):
        ai = call_llm(build_prompt(overview, exec_summary))

    schedule = build_schedule(ai["tasks"], start_date, end_date)

    with tempfile.TemporaryDirectory() as tmp:
        docx_path = os.path.join(tmp, "report.docx")
        generate_word(ai, schedule, docx_path)
        doc_bytes = open(docx_path, "rb").read()

    st.success("Report Generated")
    st.download_button("Download Word", doc_bytes, "Project_Report.docx")
